# conda activate selenium
from selenium import webdriver
from time import sleep
import re
import json
import sys
import os
from docx import Document

d = webdriver.Chrome()
pros = {}


class Problem:
    def __init__(self, problem, selections, correct):
        self.problem = problem
        self.selections = selections
        self.correct = correct

    @classmethod
    def from_dic(self, dic):
        return self(dic["problem"], dic["selections"], dic["correct"])

    def __eq__(self, other):
        return (
            self.problem == other.problem
            and self.selections == other.selections
            and self.correct == other.correct
        )

    def get_ans(self):
        return self.correct.split("、")

    def get_ans_by_current_selections(self, selections):
        ans = self.get_ans_content()
        ret = ""
        for k, v in selections.items():
            if v in ans:
                ret += k
        if len(ret) == len(self.correct):
            return ret
        else:
            return self.correct

    def get_ans_content(self):
        ret = []
        for e in self.get_ans():
            ret.append(self.selections[e])
        return ret

    def to_readable(self):
        return {
            "problem": self.problem,
            "selections": self.selections,
            "correct": self.correct,
        }


def regex_str(str):
    return str


def waiting_for_url(url):
    while d.current_url != url:
        pass


def update_problem(problem, selections, correct):
    p = Problem(problem, selections, correct)
    if not problem in pros:
        if problem != "无题干数据":
            pros[problem] = p
            print("获取到")
            print(p.to_readable())
        else:
            pros[problem] = [p]
            print("特殊更新")
            print("获取到")
            print(p.to_readable())
    else:
        print("存在相同的，进行对比")
        if problem == "无题干数据":
            print("特殊更新")
            tag = 1
            for e in pros[problem]:
                if e == p:
                    print("相同，无需更新")
                    tag = 0
            if tag:
                print("不同更新为")
                pros[problem].append(p)
                print(p.to_readable())
        else:
            if pros[problem] == p:
                print("相同，无需更新")
            else:
                print("不同更新为")
                print(p.to_readable())
                pros[problem] = p


def download_memory(url):
    d.get(url)
    sleep(1)
    while True:
        content = regex_str(
            "\n".join(
                d.find_elements_by_tag_name("main")[0].text.split("\n")[2::][:-3:]
            )
        ).split("\n")
        problem = content[0]
        selections = {}
        for e in content[1:-1:]:
            if e[0].isalpha():
                selections[e[0]] = e[1::]
            else:
                problem += "\n"
                problem += e
        if content[-1].startswith("正确答案"):
            correct = content[-1][5::]
        else:
            print(content)
            print("[Error]这个题目有点问题，没有记录")
            sleep(5)
            return
        update_problem(problem, selections, correct)
        if not d.find_elements_by_tag_name("button")[1].is_enabled():
            break
        d.find_elements_by_tag_name("button")[1].click()
        sleep(1)


def save_to_text():
    f = open(file, "w")
    ret = {}
    for k, v in pros.items():
        if k != "无题干数据":
            ret[k] = v.to_readable()
        else:
            ret[k] = []
            for e in v:
                ret[k].append(e.to_readable())
    f.write(json.dumps(ret))
    f.close()


def read_from_file():
    try:
        f = open(file, "r")
        ret = f.read()
        for k, v in json.loads(ret).items():
            if k != "无题干数据":
                pros[k] = Problem.from_dic(v)
            else:
                pros[k] = []
                for e in v:
                    pros[k].append(Problem.from_dic(e))
    except:
        print(f"读取{file}可能遇到了问题，不如删掉试试吧")


def make_result():
    while 1:
        content = regex_str(
            "\n".join(
                d.find_elements_by_tag_name("main")[0].text.split("\n")[2::][:-3:]
            )
        ).split("\n")[:-1:]
        problem = content[0]
        selections = {}
        for e in content[1::]:
            if e[0].isalpha():
                selections[e[0]] = e[1::]
            else:
                problem += "\n"
                problem += e
        ans = []
        if problem in pros:
            if problem != "无题干数据":
                p = pros[problem]
                ans = p.get_ans_by_current_selections(selections)
                print(f"找到“{problem}”\n对应答案{ans}")
            else:
                for e in pros[problem]:
                    if e.selections == selections:
                        ans = e.get_ans()
        for e in ans:
            for i in d.find_elements_by_tag_name("li"):
                if i.text.startswith(e):
                    try:
                        i.click()
                    except:
                        d.execute_script(
                            "window.scrollTo(0, document.body.scrollHeight)"
                        )
                        i.click()
        if len(ans) == 0:
            print(f"没找到“{problem}”")
            print("请自行选择答案，不用按下下一题")
            os.system("pause>nul")
        sleep(1)
        if not d.find_elements_by_tag_name("button")[2].is_enabled():
            break
        d.find_elements_by_tag_name("button")[2].click()
    sleep(1)
    d.find_elements_by_tag_name("button")[0].click()
    sleep(1)
    d.find_elements_by_tag_name("button")[2].click()


def get_len():
    ret = len(pros.items())
    if "无题干数据" in pros:
        ret += len(pros["无题干数据"]) - 1
    return ret


def add_problem(doc):
    doc.add_heading("行测题库 by 恒星", 0)
    para = doc.add_paragraph()
    i = 1

    def func(p: Problem):
        nonlocal i
        problem = p.problem
        selections = p.selections
        correct = p.correct
        para.add_run(f"{i}.{problem}").bold = True
        para.add_run().add_break()
        for k, v in selections.items():
            para.add_run(f"    {k}{v}")
            para.add_run().add_break()
        para.add_run(f"正确答案{correct}").italic = True
        para.add_run().add_break()
        i += 1

    return func


file = "data.txt"
exam = "https://exam.yooc.me/group/9755188/exams"
reviews = ["https://exam.yooc.me/group/9755112/exam/490501/review/134408164"]


def login():
    read_from_file()
    d.get(exam)
    waiting_for_url("https://www.yooc.me/mobile/yooc")
    d.get(exam)


def train():
    """
    for e in reviews:
        download_memory(e)
    """
    print("请进入考试界面")
    os.system("pause>nul")
    sleep(1)
    make_result()
    sleep(10)
    d.find_elements_by_tag_name("button")[-1].click()
    sleep(1)
    download_memory(d.current_url)
    print(f"本次更新完后题目数量为{get_len()}")
    save_to_text()
    # d.close()


def train_by_file_2025_5_6():
    def split_by_pattern(text):
        """
        按照"A.字符串"的格式分割字符串
        """
        # 使用正向肯定预查匹配字母+点号的模式作为分割点
        pattern = r"(?=[A-Za-z]\.)"
        # 进行分割并过滤可能的空字符串
        result = [s for s in re.split(pattern, text) if s]
        return result

    read_from_file()
    f = open("./2025.5.6.txt", encoding="utf-8")
    data = f.read()
    f.close()
    pattern = r"(\d+、)\s*([\s\S]*?)(A\..*?)(正确答案：[\s\S]*?)(?=\d+、|$)"
    question_blocks = re.findall(pattern, data, re.DOTALL)
    for question in question_blocks:
        problem = question[1].replace("\n", "")
        correct = question[3].replace("\n", "")[5::]
        selections = {}
        for per in split_by_pattern(question[2].replace("\n", "")):
            try:
                if per == "":
                    continue
                temp = per.split(".")
                selections[temp[0]] = f".{temp[1]}"
            except:
                print(f"出现问题{question}")
                raise
        update_problem(problem, selections, correct)
    save_to_text()


def word():
    read_from_file()
    doc = Document()
    writer = add_problem(doc)
    for k, v in pros.items():
        if k != "无题干数据":
            writer(v)
        else:
            for e in v:
                writer(e)
    doc.save("result.docx")


def main():
    print("请进入考试后按下任意键继续")
    os.system("pause>nul")
    make_result()
    print("按下任意键退出程序")
    os.system("pause>nul")


if __name__ == "__main__":
    if sys.argv[1] == "train":
        login()
        train()
    elif sys.argv[1] == "word":
        word()
    else:
        login()
        main()
    d.close()

# the logging finished
